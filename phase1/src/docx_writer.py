"""Convert a list of formatted Paragraphs into a .docx file."""

from __future__ import annotations

from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .formatting import Paragraph


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def write_docx(paragraphs: Sequence[Paragraph], output_path: str) -> str:
    """Render ``paragraphs`` to a Word document at ``output_path``.

    Each paragraph becomes one Word paragraph with the detected alignment.
    Each word becomes its own ``Run`` so per-word bold/italic flags survive.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for para in paragraphs:
        if not para.lines:
            continue

        p = doc.add_paragraph()
        p.alignment = _ALIGN_MAP.get(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        first_word_in_para = True
        for line_idx, line in enumerate(para.lines):
            for wf in line.words:
                if not first_word_in_para:
                    p.add_run(" ")
                first_word_in_para = False
                run = p.add_run(wf.word.text)
                if wf.bold:
                    run.bold = True
                if wf.italic:
                    run.italic = True
            # Preserve hard line breaks within a paragraph (e.g. addresses).
            if line_idx != len(para.lines) - 1:
                p.add_run().add_break()

    doc.save(output_path)
    return output_path
