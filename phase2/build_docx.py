"""Convert phase2/Phase2_Document.md into a polished Phase2_Document.docx.

This is a small, dedicated Markdown -> DOCX converter that handles exactly the
subset used by Phase2_Document.md:

  * # / ## / ### headings
  * Markdown pipe tables with header row + separator
  * `- ` bullet lists, `1. ` numbered lists, nested via leading spaces
  * **bold**, *italic*, `inline code`
  * ``` fenced ``` code blocks (rendered as monospace, preserved layout)
  * > block quotes
  * --- horizontal rules
  * paragraphs (blank-line separated)

We deliberately avoid pulling in pandoc / markdown / docutils — keeps the
dependency surface zero-extra over Phase 1.

Run from the repo root:
    python phase2/build_docx.py
"""

from __future__ import annotations

import os
import re
import sys
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "Phase2_Document.md")
OUT = os.path.join(HERE, "Phase2_Document.docx")


# --------------------------------------------------------------------------
# Inline span parser
# --------------------------------------------------------------------------

# Order matters: we want the longest delimiters tried first.
_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+?\*\*"        # **bold**
    r"|\*[^*\n]+?\*"             # *italic*
    r"|_[^_\n]+?_"               # _italic_
    r"|`[^`\n]+?`"               # `code`
    r"|\[[^\]]+?\]\([^)]+?\))",  # [text](url)
)


def _emit_runs(paragraph, text: str, *, base_bold=False, base_italic=False) -> None:
    """Add styled runs to `paragraph` from a markdown-formatted string."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()]).bold = base_bold or False
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            r.italic = base_italic
        elif token.startswith("*") or token.startswith("_"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            r.bold = base_bold
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        elif token.startswith("["):
            link_text = re.match(r"\[([^\]]+)\]", token).group(1)
            r = paragraph.add_run(link_text)
            r.font.color.rgb = RGBColor(0x1A, 0x4D, 0xA8)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        tail = paragraph.add_run(text[pos:])
        tail.bold = base_bold or False
        tail.italic = base_italic or False


def _set_cell_shading(cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


# --------------------------------------------------------------------------
# Block parser
# --------------------------------------------------------------------------

def _is_table_separator(line: str) -> bool:
    """A markdown table separator is `| --- | :--- | ---: | ... |`."""
    if not line.lstrip().startswith("|"):
        return False
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells) and len(cells) >= 1


def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BFBFBF")
    pbdr.append(bot)
    p_pr.append(pbdr)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j in range(n_cols):
            cell = cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _emit_runs(p, text, base_bold=(i == 0))
            if i == 0:
                _set_cell_shading(cell, "1F3B73")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True


def _add_code_block(doc: Document, lines: Iterable[str]) -> None:
    body = "\n".join(lines)
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    p_pr.append(shd)
    run = p.add_run(body)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_block_quote(doc: Document, lines: list[str]) -> None:
    text = " ".join(line.lstrip("> ").rstrip() for line in lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    _emit_runs(p, text, base_italic=True)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), "1F3B73")
    pbdr.append(left)
    p_pr.append(pbdr)


def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=min(level, 4))
    _emit_runs(p, text)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)


def _list_item(doc: Document, text: str, *, ordered: bool, depth: int) -> None:
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * depth)
    _emit_runs(p, text)


def convert(md_path: str, out_path: str) -> None:
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Body style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # blank line
        if not stripped.strip():
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            _heading(doc, m.group(2), level)
            i += 1
            continue

        # fenced code
        if stripped.lstrip().startswith("```"):
            i += 1
            buf: list[str] = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            _add_code_block(doc, buf)
            continue

        # block quote
        if stripped.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(lines[i])
                i += 1
            _add_block_quote(doc, buf)
            continue

        # table: header row + separator + body
        if (stripped.lstrip().startswith("|")
                and i + 1 < n and _is_table_separator(lines[i + 1])):
            header = _split_row(stripped)
            i += 2
            rows: list[list[str]] = [header]
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            _add_table(doc, rows)
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)", stripped)
        if m:
            depth = len(m.group(1)) // 2
            _list_item(doc, m.group(2), ordered=False, depth=depth)
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", stripped)
        if m:
            depth = len(m.group(1)) // 2
            _list_item(doc, m.group(2), ordered=True, depth=depth)
            i += 1
            continue

        # regular paragraph (concatenate consecutive non-blank, non-special lines)
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^(#{1,4})\s+", nxt):
                break
            if nxt.lstrip().startswith(("```", ">", "|", "- ", "* ")):
                break
            if re.match(r"^\s*\d+\.\s+", nxt):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        _emit_runs(p, para_text)

    doc.save(out_path)


def main() -> int:
    if not os.path.isfile(SRC):
        print(f"ERROR: source not found: {SRC}", file=sys.stderr)
        return 1
    convert(SRC, OUT)
    size_kb = os.path.getsize(OUT) // 1024
    print(f"Wrote {OUT} ({size_kb} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
