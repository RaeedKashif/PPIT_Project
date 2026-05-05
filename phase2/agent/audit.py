"""Structured audit logging + per-document audit page generator (slide 32).

Two responsibilities:
  1. Write one `audit/<job>.json` per processed job. PII is scrubbed
     (CNICs, dates, email addresses) so an attacker who reads the audit
     directory cannot reconstruct sensitive content.
  2. Append a one-page "audit page" to the produced .docx that summarizes
     what the agent did, with low-confidence words highlighted.
"""

from __future__ import annotations

import json
import os
import re
import time
from dataclasses import asdict, is_dataclass
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt


# -- PII scrubbing ----------------------------------------------------------

_CNIC_RE = re.compile(r"\b\d{5}-?\d{7}-?\d\b")           # Pakistan CNIC pattern
_EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b")
_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b"
)
_PHONE_RE = re.compile(r"\b\+?\d{2,4}[-\s]?\d{3,4}[-\s]?\d{4}\b")


def scrub_pii(text: str) -> str:
    """Replace detected PII patterns with type-tagged placeholders."""
    if not text:
        return text
    text = _CNIC_RE.sub("[CNIC]", text)
    text = _EMAIL_RE.sub("[EMAIL]", text)
    text = _PHONE_RE.sub("[PHONE]", text)
    text = _DATE_RE.sub("[DATE]", text)
    return text


def _to_jsonable(obj: Any) -> Any:
    if is_dataclass(obj):
        return _to_jsonable(asdict(obj))
    if isinstance(obj, dict):
        return {k: _to_jsonable(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_to_jsonable(v) for v in obj]
    if isinstance(obj, str):
        return scrub_pii(obj)
    if isinstance(obj, (int, float, bool)) or obj is None:
        return obj
    return str(obj)


# -- structured event log ---------------------------------------------------

class AuditLogger:
    def __init__(self, audit_dir: str, *, max_files: int = 1000):
        self.audit_dir = audit_dir
        self.max_files = max_files
        os.makedirs(self.audit_dir, exist_ok=True)

    def write(self, event: dict[str, Any]) -> str:
        job_id = event.get("job_id") or time.strftime("%Y-%m-%dT%H-%M-%S")
        # sanitize for filename
        safe = re.sub(r"[^A-Za-z0-9_\-]", "_", str(job_id))
        path = os.path.join(self.audit_dir, f"{safe}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(_to_jsonable(event), f, indent=2)
        self._rotate()
        return path

    def _rotate(self) -> None:
        files = [
            os.path.join(self.audit_dir, n)
            for n in os.listdir(self.audit_dir)
            if n.endswith(".json")
        ]
        if len(files) <= self.max_files:
            return
        files.sort(key=os.path.getmtime)
        for old in files[: len(files) - self.max_files]:
            try:
                os.remove(old)
            except OSError:
                pass


# -- audit page appended to the .docx ---------------------------------------

def append_audit_page(
    docx_path: str,
    *,
    job_id: str,
    doc_type: str,
    plan_summary: dict[str, Any],
    aggregate_confidence: float,
    word_count: int,
    low_conf_words: list[tuple[str, float]],
    fired_rules: list[str] | None = None,
    llm_suggestions: list[dict[str, Any]] | None = None,
) -> None:
    """Append a structured 'what the agent did' page to the produced doc."""
    doc = Document(docx_path)
    doc.add_page_break()

    h = doc.add_paragraph()
    r = h.add_run("DocAgent — Audit Page")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(f"Job ID: {job_id}")
    doc.add_paragraph(f"Document type: {doc_type}")
    doc.add_paragraph(
        f"Plan: preprocessor={plan_summary.get('preprocessor')}, "
        f"psm={plan_summary.get('psm')}, lang={plan_summary.get('lang')}, "
        f"spell_check={plan_summary.get('use_spell_check')}, "
        f"llm_repair={plan_summary.get('use_llm_repair')}"
    )
    doc.add_paragraph(
        f"Aggregate confidence: {aggregate_confidence:.2f}  "
        f"(words: {word_count}, low-confidence: {len(low_conf_words)})"
    )

    if fired_rules:
        doc.add_paragraph("Memory rules that fired this run:")
        for rule in fired_rules:
            doc.add_paragraph(f"  • {scrub_pii(rule)}")

    if llm_suggestions:
        doc.add_paragraph(
            "LLM-suggested corrections (each requires user approval before "
            "being written above):"
        )
        for s in llm_suggestions:
            line = doc.add_paragraph(
                f"  • {scrub_pii(s.get('original', ''))} → "
                f"{scrub_pii(s.get('suggested', ''))} "
                f"[{s.get('status', 'pending')}]"
            )
            line.runs[0].italic = True

    if low_conf_words:
        doc.add_paragraph("Low-confidence words (highlighted yellow above):")
        para = doc.add_paragraph()
        for i, (word, conf) in enumerate(low_conf_words[:50]):
            run = para.add_run(f"{scrub_pii(word)}({conf:.2f})")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if i != len(low_conf_words[:50]) - 1:
                para.add_run("  ")

    doc.add_paragraph(
        "This page exists so you can audit what the agent did. PII patterns "
        "(CNICs, emails, phone numbers, dates) are masked here for safety."
    )

    doc.save(docx_path)
